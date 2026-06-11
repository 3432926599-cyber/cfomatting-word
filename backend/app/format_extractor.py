"""
格式提取引擎 —— 从上传的 .docx 范本中提取格式规则。
使用 python-docx 高层 API + lxml 底层 OXML 直读兜底。
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import json
import re
from typing import Optional, Any


# ── 对齐方式映射 ─────────────────────────────────────────
ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def _pt_or_none(emu) -> Optional[float]:
    """Emu → pt，失败返回 None"""
    try:
        if emu is None:
            return None
        return round(emu / 12700, 1)  # 1pt = 12700 EMU
    except Exception:
        return None


def _cm_or_none(emu) -> Optional[str]:
    """Emu → cm 字符串"""
    try:
        if emu is None:
            return None
        cm = emu / 360000
        return f"{round(cm, 2)}cm"
    except Exception:
        return None


def _resolve_run_font(run, doc) -> dict:
    """
    解析单个 run 的字体属性。
    若 python-docx 返回 None → 尝试从底层 XML 读取 → 再兜底返回 None。
    """
    font = run.font
    result = {}

    # 字体名称
    name = font.name
    if name is None:
        # 尝试从 rPr/rFonts 读取
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                name = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:hAnsi"))
    result["font_name"] = name

    # 字号
    size = font.size
    if size is None:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                size = Pt(int(sz.get(qn("w:val"))) / 2)
    result["font_size"] = _pt_or_none(size) if size else None

    # 加粗
    bold = font.bold
    if bold is None:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            b = rPr.find(qn("w:b"))
            bold = b is not None
    result["bold"] = bool(bold) if bold is not None else None

    # 斜体
    italic = font.italic
    if italic is None:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            i = rPr.find(qn("w:i"))
            italic = i is not None
    result["italic"] = bool(italic) if italic is not None else None

    # 颜色
    try:
        if font.color and font.color.rgb:
            result["color"] = str(font.color.rgb)
    except Exception:
        result["color"] = None

    return {k: v for k, v in result.items() if v is not None}


def _resolve_paragraph_format(para) -> dict:
    """解析段落格式属性"""
    pf = para.paragraph_format
    result = {}

    # 对齐方式
    if pf.alignment is not None:
        result["alignment"] = ALIGN_MAP.get(pf.alignment, "left")

    # 首行缩进
    if pf.first_line_indent is not None:
        result["first_line_indent"] = _cm_or_none(pf.first_line_indent)

    # 左缩进
    if pf.left_indent is not None:
        result["left_indent"] = _cm_or_none(pf.left_indent)

    # 行距（返回倍数或固定值）
    if pf.line_spacing is not None:
        result["line_spacing"] = round(pf.line_spacing, 2)

    # 行距规则
    if pf.line_spacing_rule is not None:
        rule_map = {
            0: "single",
            1: "double",
            2: "multiple",
            3: "exactly",
            4: "at_least",
        }
        result["line_spacing_rule"] = rule_map.get(pf.line_spacing_rule, "multiple")

    # 段前 / 段后
    if pf.space_before is not None:
        result["space_before"] = _cm_or_none(pf.space_before)
    if pf.space_after is not None:
        result["space_after"] = _cm_or_none(pf.space_after)

    # 大纲级别（python-docx 某些版本无此属性，从底层 XML 读取）
    try:
        if pf.outline_level is not None:
            result["outline_level"] = pf.outline_level
    except AttributeError:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                result["outline_level"] = int(ol.get(qn("w:val")))

    return result


def _get_style_properties(doc, style_name: str) -> Optional[dict]:
    """
    沿样式继承链追溯，返回该样式最终生效的字体+段落属性。
    通过遍历使用该样式的段落来获取实际属性。
    """
    font_attrs = {}
    para_attrs = {}

    # 查找使用该样式且有文本的段落
    sample_para = None
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name and para.text.strip():
            sample_para = para
            break

    if sample_para is None:
        return None

    # 提取字体属性（取第一个 run）
    if sample_para.runs:
        font_attrs = _resolve_run_font(sample_para.runs[0], doc)

    # 提取段落属性
    para_attrs = _resolve_paragraph_format(sample_para)

    return {"font": font_attrs, "paragraph": para_attrs}


def _get_section_properties(doc) -> dict:
    """提取第一个节的页面设置"""
    result = {}
    if doc.sections:
        sec = doc.sections[0]
        result["page_width"] = _cm_or_none(sec.page_width)
        result["page_height"] = _cm_or_none(sec.page_height)
        result["margin_top"] = _cm_or_none(sec.top_margin)
        result["margin_bottom"] = _cm_or_none(sec.bottom_margin)
        result["margin_left"] = _cm_or_none(sec.left_margin)
        result["margin_right"] = _cm_or_none(sec.right_margin)
    return result


def _get_header_footer(doc) -> dict:
    """提取页眉页脚信息"""
    result = {}
    if doc.sections:
        sec = doc.sections[0]

        # 页眉
        try:
            header = sec.header
            if header and not header.is_linked_to_previous:
                text = "".join(p.text for p in header.paragraphs)
                if text.strip():
                    result["header_text"] = text.strip()
        except Exception:
            pass

        # 页脚
        try:
            footer = sec.footer
            if footer and not footer.is_linked_to_previous:
                text = "".join(p.text for p in footer.paragraphs)
                if text.strip():
                    result["footer_text"] = text.strip()
                # 检测是否有页码
                for p in footer.paragraphs:
                    xml = etree.tostring(p._element, encoding="unicode")
                    if "PAGE" in xml or "page" in xml.lower():
                        result["has_page_number"] = True
                        break
        except Exception:
            pass

        # 奇偶页差异
        try:
            sectPr = sec._sectPr
            if sectPr is not None:
                titlePg = sectPr.find(qn("w:titlePg"))
                result["different_first_page"] = titlePg is not None
                evenAndOddHeaders = sectPr.find(qn("w:evenAndOddHeaders"))
                result["different_odd_even"] = evenAndOddHeaders is not None
        except Exception:
            pass

    return result


def _extract_by_actual_formatting(doc) -> dict:
    """
    按段落实际格式聚类识别 —— 适用于"直接格式"的范本
    （老师选中文字手动调字体加粗，不依赖Word样式名）。

    策略：
    1. 遍历所有段落，提取每条的实际字体/字号/加粗/对齐/行距
    2. 按 (字体名, 字号, 加粗) 聚类
    3. 字号最大的 → Title，次大加粗的 → Heading1，依次类推
    4. 最常见且字数多的 → Normal（正文）
    """
    from collections import defaultdict

    # ── 收集每个段落的格式指纹 ──
    groups = defaultdict(list)  # key: (font_name, font_size, bold) → [paragraphs]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not para.runs:
            continue

        run = para.runs[0]
        font_info = _resolve_run_font(run, doc)
        para_info = _resolve_paragraph_format(para)

        fn = font_info.get("font_name", "?")
        fs = font_info.get("font_size", 0)
        bd = font_info.get("bold", False)

        if fs == 0:
            continue

        key = (fn, round(fs), bd)
        groups[key].append({
            "text": text,
            "len": len(text),
            "font": font_info,
            "paragraph": para_info,
        })

    if not groups:
        return {}

    # ── 按字号从大到小排序各组 ──
    sorted_groups = sorted(groups.items(), key=lambda x: -x[0][1])

    # ── 识别正文组：总字数最多的组 ──
    best_body = None
    best_body_chars = 0
    for key, items in groups.items():
        total_chars = sum(it["len"] for it in items)
        if total_chars > best_body_chars:
            best_body_chars = total_chars
            best_body = key

    # ── 分配样式 ──
    styles = {}
    heading_level = 0  # 0=Title, 1=H1, 2=H2, ...

    for key, items in sorted_groups:
        fn, fs, bd = key
        sample = items[0]

        if key == best_body:
            # 这是正文
            styles["Normal"] = {
                "font": sample["font"],
                "paragraph": sample["paragraph"],
            }
        elif bd and fs > best_body[1]:
            # 加粗 + 比正文大 → 标题层级
            label = "Title" if heading_level == 0 else f"Heading{heading_level}"
            styles[label] = {
                "font": sample["font"],
                "paragraph": sample["paragraph"],
            }
            heading_level += 1
        elif not bd and fs > best_body[1]:
            # 不加粗但比正文大 → 可能是特殊格式（如封面大字）
            if heading_level == 0:
                styles["Title"] = {
                    "font": sample["font"],
                    "paragraph": sample["paragraph"],
                }
                heading_level += 1

    # ── 如果Title跟Heading1一样大，合并 ──
    if "Title" in styles and "Heading1" in styles:
        t_size = styles["Title"]["font"].get("font_size", 0)
        h1_size = styles["Heading1"]["font"].get("font_size", 0)
        if abs(t_size - h1_size) < 1:
            styles["Heading1"] = styles.pop("Title")

    return styles


def extract_format(docx_path: str) -> dict:
    """
    主入口 —— 从 .docx 文件提取完整格式规则。

    返回格式：
    {
        "page_setup": { ... },
        "styles": {
            "Title":   { "font": {...}, "paragraph": {...} },
            "Heading1": { ... },
            "Heading2": { ... },
            "Heading3": { ... },
            "Normal":  { ... }
        },
        "header_footer": { ... }
    }
    """
    doc = Document(docx_path)

    result = {
        "source": docx_path,
        "page_setup": _get_section_properties(doc),
        "styles": {},
        "header_footer": _get_header_footer(doc),
    }

    # ── 方法1（优先）：按实际格式聚类识别 ──
    result["styles"] = _extract_by_actual_formatting(doc)

    # ── 方法2（回退）：聚类结果不足时，按样式名提取 ──
    if len(result["styles"]) < 2:
        style_candidates = [
            "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4",
            "Normal", "标题", "标题 1", "标题 2", "标题 3", "标题 4",
            "正文", "正文文本",
        ]
        found_styles = {}
        for para in doc.paragraphs:
            if para.style is None:
                continue
            sname = para.style.name
            if sname not in found_styles and para.text.strip():
                found_styles[sname] = True

        priority_map = [
            ("Title", ["Title", "标题"]),
            ("Heading1", ["Heading 1", "标题 1", "heading 1"]),
            ("Heading2", ["Heading 2", "标题 2", "heading 2"]),
            ("Heading3", ["Heading 3", "标题 3", "heading 3"]),
            ("Heading4", ["Heading 4", "标题 4", "heading 4"]),
            ("Normal", ["Normal", "正文", "正文文本", "Body Text"]),
        ]

        for target_key, names in priority_map:
            for name in names:
                if name in found_styles:
                    props = _get_style_properties(doc, name)
                    if props:
                        result["styles"][target_key] = props
                    break

    # ── 兜底：至少保证 Normal 有值 ──
    if "Normal" not in result["styles"]:
        normal_font = {"font_name": "宋体", "font_size": 12.0}
        normal_para = {"line_spacing": 1.5, "alignment": "left"}
        for para in doc.paragraphs:
            if para.style and para.style.name == "Normal" and para.text.strip():
                if para.runs:
                    f = _resolve_run_font(para.runs[0], doc)
                    normal_font.update({k: v for k, v in f.items() if v is not None})
                p = _resolve_paragraph_format(para)
                normal_para.update({k: v for k, v in p.items() if v is not None})
        result["styles"]["Normal"] = {"font": normal_font, "paragraph": normal_para}

    # ── 清理：移除空值 ──
    def clean_dict(d):
        return {k: v for k, v in d.items() if v is not None and v != {}}

    result["styles"] = {k: clean_dict(v) for k, v in result["styles"].items() if v}
    result["page_setup"] = clean_dict(result["page_setup"])
    result["header_footer"] = clean_dict(result["header_footer"])

    return result
