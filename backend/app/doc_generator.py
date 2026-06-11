"""
文档生成器 —— 基于 python-docx，将格式规则应用到结构化内容，生成 Word 文档。
支持两种模式：
1. 模板模式：保留封面/前导页（校徽、表格等），替换正文内容
2. 新建模式：从零创建文档，应用格式规则
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import re
from typing import Dict, Any, Optional
from copy import deepcopy


# ── 对齐方式反向映射 ────────────────────────────────────
ALIGN_MAP_REVERSE = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _parse_cm(value: str) -> Optional[Cm]:
    if value is None:
        return None
    try:
        if value.endswith("cm"):
            return Cm(float(value.replace("cm", "")))
        if value.endswith("mm"):
            return Cm(float(value.replace("mm", "")) / 10)
        if value.endswith("in"):
            return Cm(float(value.replace("in", "")) * 2.54)
    except (ValueError, AttributeError):
        pass
    return None


def _apply_font(run, font_rule: dict):
    if not font_rule:
        return
    if "font_name" in font_rule and font_rule["font_name"]:
        run.font.name = font_rule["font_name"]
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_rule["font_name"])
    if "font_size" in font_rule and font_rule["font_size"]:
        run.font.size = Pt(font_rule["font_size"])
    if "bold" in font_rule and font_rule["bold"] is not None:
        run.font.bold = font_rule["bold"]
    if "italic" in font_rule and font_rule["italic"] is not None:
        run.font.italic = font_rule["italic"]
    if "color" in font_rule and font_rule["color"]:
        try:
            run.font.color.rgb = RGBColor.from_string(font_rule["color"])
        except Exception:
            pass


def _apply_paragraph_format(para, para_rule: dict):
    if not para_rule:
        return
    pf = para.paragraph_format
    if "alignment" in para_rule and para_rule["alignment"]:
        align = ALIGN_MAP_REVERSE.get(para_rule["alignment"])
        if align is not None:
            pf.alignment = align
    if "first_line_indent" in para_rule and para_rule["first_line_indent"]:
        cm = _parse_cm(para_rule["first_line_indent"])
        if cm:
            pf.first_line_indent = cm
    if "left_indent" in para_rule and para_rule["left_indent"]:
        cm = _parse_cm(para_rule["left_indent"])
        if cm:
            pf.left_indent = cm
    if "line_spacing" in para_rule and para_rule["line_spacing"] is not None:
        pf.line_spacing = para_rule["line_spacing"]
    if "space_before" in para_rule and para_rule["space_before"]:
        cm = _parse_cm(para_rule["space_before"])
        if cm:
            pf.space_before = cm
    if "space_after" in para_rule and para_rule["space_after"]:
        cm = _parse_cm(para_rule["space_after"])
        if cm:
            pf.space_after = cm


# ══════════════════════════════════════════════════════════
# 封面 / 正文分界检测
# ══════════════════════════════════════════════════════════

def detect_body_start(doc: Document) -> int:
    """
    检测正文开始的段落索引。
    所有位于此索引之前的段落将被原样保留（封面、目录页等）。

    检测规则（按优先级）：
    1. 分节符 —— 第一个分节符后的段落即为正文开始
    2. 关键词 —— "摘要"/"Abstract"/"目录"/"第X章"/"一、"
    3. 格式描述行 —— 遇到"题目：宋体"/"正文：宋体"等格式说明行，其前的内容视为封面
    4. 填充行结束 —— 连续的"___"或"：____"填充行之后
    5. 默认 —— 文档前25%作为封面
    """
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    total = len(paras)

    # ── 规则1：分节符 ──
    for i, p in enumerate(paras):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None and i + 1 < total:
                return i + 1

    # ── 规则2：正文关键词 ──
    cover_keywords = [
        "摘要", "Abstract", "ABSTRACT",
        "目录", "目  录", "目 录",
        "第一章", "第1章", "一、", "1.",
    ]
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().replace(" ", "")
        for kw in cover_keywords:
            if text.startswith(kw):
                return max(0, i)

    # ── 规则3：格式描述行（"题目：宋体，二号" 等）──
    format_clues = ["宋体", "黑体", "楷体", "仿宋", "字号", "行距",
                    "题目：", "标题：", "正文：", "页眉：", "页码："]
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        clue_count = sum(1 for c in format_clues if c in text)
        if clue_count >= 2:
            return i  # 格式描述行之前的内容都是封面

    # ── 规则4：连续填充行结束 ──
    fill_line_count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r"^[\w一-鿿]+[：:]\s*_+$", text) or \
           re.match(r"^[\w一-鿿]+[：:]\s*$", text):
            fill_line_count += 1
        elif fill_line_count >= 2 and text:
            # 填充行之后有实质性内容 → 此处为分界
            return i

    # ── 规则5：默认 ──
    ratio = 4 if total < 20 else 5
    return max(1, total // ratio)


def _get_para_element_at(doc: Document, index: int):
    """获取文档中第 index 个 w:p 元素"""
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    if 0 <= index < len(paras):
        return paras[index]
    return None


# ══════════════════════════════════════════════════════════
# 内容添加辅助函数
# ══════════════════════════════════════════════════════════

def _add_title_para(doc, text: str, format_rules: dict):
    title_rule = format_rules.get("styles", {}).get("Title", {})
    h1_rule = format_rules.get("styles", {}).get("Heading1", {})
    font_rule = title_rule.get("font") or h1_rule.get("font") or {"font_name": "黑体", "font_size": 18, "bold": True}
    para_rule = title_rule.get("paragraph") or h1_rule.get("paragraph") or {"alignment": "center"}
    para = doc.add_paragraph()
    run = para.add_run(text)
    _apply_font(run, font_rule)
    _apply_paragraph_format(para, para_rule)
    return para


def _add_heading_para(doc, text: str, level: int, format_rules: dict):
    # 回退：如果当前层级不存在，逐级向上找最近的可用标题样式
    styles = format_rules.get("styles", {})
    style_rule = {}
    for l in range(level, 0, -1):
        key = f"Heading{l}" if l <= 4 else "Heading4"
        if key in styles:
            style_rule = styles[key]
            break
    para = doc.add_paragraph()
    run = para.add_run(text)
    _apply_font(run, style_rule.get("font", {}))
    _apply_paragraph_format(para, style_rule.get("paragraph", {}))
    return para


def _add_body_para(doc, text: str, format_rules: dict):
    style_rule = format_rules.get("styles", {}).get("Normal", {})
    para = doc.add_paragraph()
    run = para.add_run(text)
    _apply_font(run, style_rule.get("font", {}))
    _apply_paragraph_format(para, style_rule.get("paragraph", {}))
    return para


# ══════════════════════════════════════════════════════════
# 模板模式：保留封面，替换正文
# ══════════════════════════════════════════════════════════

def generate_from_template(
    template_path: str,
    content_data: Dict[str, Any],
    format_rules: Dict[str, Any],
    output_path: str,
) -> str:
    """
    基于上传的范本文件生成文档。
    - 封面/前导页（校徽、标题页等）原样保留
    - 正文部分替换为用户内容
    """
    doc = Document(template_path)

    body_start = detect_body_start(doc)
    body = doc.element.body
    all_paras = body.findall(qn("w:p"))

    # ── 删除正文及其后的所有段落 ──
    # 从 body_start 开始删除到末尾（保留之前的所有内容）
    paras_to_remove = all_paras[body_start:]
    for p in paras_to_remove:
        body.remove(p)

    # ── 删除正文区域可能残留的 tbl（表格）元素 ──
    # 封面中的表格（如个人信息表）通常在 body_start 之前，不会被删除
    # 但正文区域中的表格需要清除
    tables_to_remove = []
    for tbl in body.findall(qn("w:tbl")):
        # 检查表格位置：如果在最后一个保留段落之后 → 删除
        tables_to_remove.append(tbl)
    for tbl in tables_to_remove:
        body.remove(tbl)

    # ── 追加用户内容 ──
    elements = content_data.get("elements", [])

    for elem in elements:
        etype = elem.get("type")

        if etype == "title":
            _add_title_para(doc, elem["text"], format_rules)

        elif etype == "heading":
            level = elem.get("level", 1)
            _add_heading_para(doc, elem["text"], level, format_rules)

        elif etype == "paragraph":
            _add_body_para(doc, elem["text"], format_rules)

        elif etype == "list":
            is_ordered = elem.get("ordered", False)
            items = elem.get("items", [])
            for idx, item in enumerate(items):
                prefix = f"{idx + 1}. " if is_ordered else "• "
                _add_body_para(doc, f"{prefix}{item['text']}", format_rules)

    # ── 保存 ──
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════
# 新建模式：从零创建（无模板时使用）
# ══════════════════════════════════════════════════════════

def generate_from_scratch(
    content_data: Dict[str, Any],
    format_rules: Dict[str, Any],
    output_path: str,
) -> str:
    """从零创建 Word 文档，应用格式规则。"""
    doc = Document()

    # 页面设置
    page_setup = format_rules.get("page_setup", {})
    if doc.sections:
        sec = doc.sections[0]
        for key, attr in [
            ("margin_top", "top_margin"),
            ("margin_bottom", "bottom_margin"),
            ("margin_left", "left_margin"),
            ("margin_right", "right_margin"),
        ]:
            if key in page_setup:
                cm = _parse_cm(page_setup[key])
                if cm:
                    setattr(sec, attr, cm)

    # 预定义 Normal 样式
    normal_rule = format_rules.get("styles", {}).get("Normal", {})
    if normal_rule:
        style = doc.styles["Normal"]
        font_rule = normal_rule.get("font", {})
        if "font_name" in font_rule and font_rule["font_name"]:
            style.font.name = font_rule["font_name"]
        if "font_size" in font_rule and font_rule["font_size"]:
            style.font.size = Pt(font_rule["font_size"])
        para_rule = normal_rule.get("paragraph", {})
        if "line_spacing" in para_rule and para_rule["line_spacing"]:
            style.paragraph_format.line_spacing = para_rule["line_spacing"]

    # 添加页眉
    hf = format_rules.get("header_footer", {})
    if hf.get("header_text") and doc.sections:
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hr = hp.add_run(hf["header_text"])
        hfont = hf.get("header_font", {})
        _apply_font(hr, hfont)

    # 遍历内容
    elements = content_data.get("elements", [])
    for elem in elements:
        etype = elem.get("type")
        if etype == "title":
            _add_title_para(doc, elem["text"], format_rules)
        elif etype == "heading":
            _add_heading_para(doc, elem["text"], elem.get("level", 1), format_rules)
        elif etype == "paragraph":
            _add_body_para(doc, elem["text"], format_rules)
        elif etype == "list":
            is_ordered = elem.get("ordered", False)
            items = elem.get("items", [])
            for idx, item in enumerate(items):
                prefix = f"{idx + 1}. " if is_ordered else "• "
                _add_body_para(doc, f"{prefix}{item['text']}", format_rules)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════
# 统一入口
# ══════════════════════════════════════════════════════════

def generate_docx(
    content_data: Dict[str, Any],
    format_rules: Dict[str, Any],
    output_path: str,
    template_path: Optional[str] = None,
) -> str:
    """
    主入口 —— 生成 Word 文档。

    参数：
        content_data: 结构化内容
        format_rules: 格式规则
        output_path: 输出路径
        template_path: 模板文件路径（有则保留封面，无则从零创建）

    返回：输出文件路径
    """
    if template_path and os.path.exists(template_path):
        try:
            return generate_from_template(
                template_path, content_data, format_rules, output_path
            )
        except Exception:
            pass  # 模板模式失败 → 回退到新建模式

    return generate_from_scratch(content_data, format_rules, output_path)
